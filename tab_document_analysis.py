# =============================================================
# tab_document_analysis.py
# 🤖 วิเคราะห์เอกสารการเงินด้วย Claude API — อัปโหลด PDF/Excel/Word (เช่น งบการเงิน, รายงาน
# ประจำปี 56-1) ให้ Claude อ่านและสรุปจุดเด่น/จุดเสี่ยงให้อัตโนมัติ
#
# ใช้ API Key ส่วนตัวของผู้ใช้เอง (ไม่ใช่ของแอป) เก็บไว้ใน Streamlit Secrets เพื่อความปลอดภัย
# ไม่ฝัง Key ไว้ในโค้ดหรือ Google Sheets โดยตรง
#
# หมายเหตุสำคัญเรื่องชนิดไฟล์:
#   - PDF: ส่งเข้า Claude API ได้ตรงๆ (Claude อ่าน PDF ได้ทั้งข้อความและภาพ/กราฟ/ตารางในตัว)
#   - Excel (.xlsx) และ Word (.docx): Claude API "ไม่รองรับ" ไฟล์ 2 ประเภทนี้โดยตรง ต้องแปลงเป็น
#     ข้อความล้วนก่อน (ใช้ pandas สำหรับ Excel, python-docx สำหรับ Word) แล้วค่อยส่งข้อความที่
#     แปลงแล้วเข้าไปแทน
# =============================================================
import streamlit as st
import pandas as pd
import base64
from backend_functions import get_active_sheet_name, save_document_analysis_history, load_document_analysis_history
import io

MODEL_NAME = "claude-sonnet-5"  # โมเดลปัจจุบันที่สมดุลระหว่างคุณภาพและราคา เหมาะกับงานสรุปเอกสาร

ANALYSIS_PROMPT = """คุณเป็นนักวิเคราะห์การเงินที่ช่วยสรุปเอกสารให้นักลงทุนรายย่อยเข้าใจง่าย
กรุณาอ่านเอกสารที่แนบมา แล้วสรุปให้ครบทั้ง 4 หัวข้อนี้ เป็นภาษาไทย กระชับ อ่านง่าย:

1. **ภาพรวมธุรกิจ** — บริษัททำธุรกิจอะไร มีรายได้หลักจากไหน
2. **จุดเด่น** — สิ่งที่น่าสนใจ/เป็นบวก เช่น การเติบโตของรายได้-กำไร, ความได้เปรียบทางธุรกิจ
3. **จุดเสี่ยง** — สัญญาณที่ควรระวัง เช่น หนี้สินสูง, กำไรลดลง, ปัจจัยเสี่ยงที่ระบุในรายงาน
4. **สรุปตัวเลขสำคัญ** — รายได้, กำไรสุทธิ, หนี้สินต่อทุน (ถ้ามีในเอกสาร) เทียบปีก่อนหน้า (ถ้ามีข้อมูล)

หมายเหตุ: นี่เป็นการสรุปข้อมูลเพื่อประกอบการตัดสินใจเท่านั้น ไม่ใช่คำแนะนำการลงทุน"""


def _extract_text_from_xlsx(file_bytes):
    """แปลงไฟล์ Excel ทุกชีตเป็นข้อความ (Markdown table) เพราะ Claude API ไม่รองรับไฟล์ .xlsx โดยตรง"""
    excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
    text_parts = []
    for sheet_name in excel_file.sheet_names:
        df = excel_file.parse(sheet_name)
        text_parts.append(f"=== ชีต: {sheet_name} ===\n{df.to_markdown(index=False)}\n")
    return "\n".join(text_parts)


def _extract_text_from_docx(file_bytes):
    """แปลงไฟล์ Word เป็นข้อความล้วน (ย่อหน้า + ตาราง) เพราะ Claude API ไม่รองรับไฟล์ .docx โดยตรง"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table_idx, table in enumerate(doc.tables):
        text_parts.append(f"\n=== ตารางที่ {table_idx + 1} ===")
        for row in table.rows:
            text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))

    return "\n".join(text_parts)


def render_tab_document_analysis():
    st.markdown("### 🤖 วิเคราะห์เอกสารการเงินด้วย AI")
    st.markdown(
        "อัปโหลดงบการเงิน, รายงานประจำปี (56-1), หรือเอกสารการเงินอื่นๆ (PDF, Excel, Word) "
        "ให้ Claude อ่านและสรุปจุดเด่น/จุดเสี่ยงให้อัตโนมัติ"
    )

    # เช็คว่ามี API Key ตั้งค่าไว้แล้วหรือยัง (เก็บใน Streamlit Secrets ไม่ใช่ในโค้ด)
    api_key = st.secrets.get("ANTHROPIC_API_KEY") if hasattr(st, "secrets") else None
    if not api_key:
        st.warning(
            "⚠️ ยังไม่ได้ตั้งค่า Claude API Key ครับ — ไปที่ Streamlit Cloud → เลือกแอปนี้ → "
            "**Settings → Secrets** แล้วเพิ่มบรรทัด:\n\n"
            "`ANTHROPIC_API_KEY = \"sk-ant-...\"`\n\n"
            "(ใช้ API Key จาก console.anthropic.com ของคุณเอง)"
        )
        return

    try:
        import anthropic
    except ImportError:
        st.error("⚠️ ยังไม่ได้ติดตั้งไลบรารี `anthropic` ครับ — เพิ่ม `anthropic` ลงใน requirements.txt แล้ว reboot แอปก่อนครับ")
        return

    uploaded_file = st.file_uploader(
        "อัปโหลดเอกสาร (PDF, Excel, Word)", type=["pdf", "xlsx", "docx"], key="doc_analysis_uploader"
    )

    if uploaded_file is None:
        st.info("💡 ตัวอย่างเอกสารที่เหมาะสม: งบการเงินรายไตรมาส, รายงานประจำปี 56-1, บทวิเคราะห์หุ้นจากโบรกเกอร์")
        return

    file_bytes = uploaded_file.read()
    file_ext = uploaded_file.name.split(".")[-1].lower()

    st.info(f"📄 ไฟล์: **{uploaded_file.name}** ({len(file_bytes) / 1024:.0f} KB)")

    if st.button("🔍 วิเคราะห์เอกสารนี้ด้วย AI", type="primary", use_container_width=True):
        with st.spinner("กำลังให้ AI อ่านและวิเคราะห์เอกสาร... (อาจใช้เวลา 10-30 วินาที ขึ้นกับความยาวเอกสาร)"):
            try:
                client = anthropic.Anthropic(api_key=api_key)

                if file_ext == "pdf":
                    # PDF ส่งเข้า Claude API ได้ตรงๆ (รองรับทั้งข้อความและภาพ/กราฟ/ตารางในตัว)
                    base64_data = base64.standard_b64encode(file_bytes).decode("utf-8")
                    message_content = [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": base64_data,
                            },
                        },
                        {"type": "text", "text": ANALYSIS_PROMPT},
                    ]
                elif file_ext == "xlsx":
                    extracted_text = _extract_text_from_xlsx(file_bytes)
                    message_content = f"เนื้อหาจากไฟล์ Excel:\n\n{extracted_text}\n\n{ANALYSIS_PROMPT}"
                elif file_ext == "docx":
                    extracted_text = _extract_text_from_docx(file_bytes)
                    message_content = f"เนื้อหาจากไฟล์ Word:\n\n{extracted_text}\n\n{ANALYSIS_PROMPT}"
                else:
                    st.error("ไม่รองรับไฟล์ประเภทนี้ครับ")
                    return

                response = client.messages.create(
                    model=MODEL_NAME,
                    max_tokens=2000,
                    messages=[{"role": "user", "content": message_content}],
                )

                result_text = "".join(block.text for block in response.content if block.type == "text")

                st.divider()
                st.markdown("#### 📋 ผลการวิเคราะห์")
                st.markdown(result_text)

                # แสดงจำนวน token ที่ใช้จริง ช่วยให้ผู้ใช้ประเมินค่าใช้จ่ายได้
                usage = response.usage
                st.caption(
                    f"💰 ใช้ token: อินพุต {usage.input_tokens:,} / เอาต์พุต {usage.output_tokens:,} "
                    f"(ราคาอ้างอิง ≈ ${usage.input_tokens/1_000_000*2 + usage.output_tokens/1_000_000*10:.4f})"
                )

                # 🆕 บันทึกผลลัพธ์ลง Google Sheets ทันที ก่อนหน้านี้แสดงแค่บนหน้าจอ พอปิด/รีเฟรช
                # หน้าเว็บ ผลลัพธ์จะหายไปเลย ไม่มีทางเรียกดูย้อนหลังได้ ตอนนี้บันทึกอัตโนมัติทุกครั้ง
                _save_success, _save_msg = save_document_analysis_history(
                    get_active_sheet_name(), uploaded_file.name, result_text
                )
                if _save_success:
                    st.caption("✅ บันทึกผลวิเคราะห์นี้ไว้แล้ว เรียกดูย้อนหลังได้ที่ด้านล่างสุดของหน้านี้")

            except Exception as e:
                st.error(f"❌ เกิดข้อผิดพลาดในการวิเคราะห์: {e}")

    # --- ประวัติผลวิเคราะห์ย้อนหลัง ---
    st.divider()
    st.markdown("#### 📜 ประวัติผลวิเคราะห์ย้อนหลัง")
    df_history = load_document_analysis_history(get_active_sheet_name())
    if df_history.empty:
        st.caption("ยังไม่มีประวัติการวิเคราะห์เลยครับ")
    else:
        for _, row in df_history.sort_values('Date', ascending=False).iterrows():
            with st.expander(f"📄 {row.get('Filename', '?')} — {row.get('Date', '?')}"):
                st.markdown(row.get('Analysis_Result', ''))
